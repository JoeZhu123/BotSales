from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
from typing import List, Dict

class ReportGenerator:
    """
    Word 报告生成器：将分析结果、图表和数据转化为专业的文档
    """
    def __init__(self, output_dir: str = "data/reports"):
        self.output_dir = output_dir
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def generate_word_report(self, keyword: str, analysis: Dict, sales_data: List[Dict], sourcing_data: List[Dict], trend_data: List[Dict], viz_path: str):
        """
        生成 Word 格式的深度分析报告
        """
        doc = Document()
        
        # 1. 标题
        title = doc.add_heading(f'全球电商选品与趋势分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2. 关键词与生成时间
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'分析关键词: {keyword}\n生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        run.font.size = Pt(12)

        # 3. 核心结论 (AI点评)
        doc.add_heading('一、 核心结论与 AI 策略建议', level=1)
        doc.add_paragraph(analysis.get('ai_analysis', '暂无 AI 分析内容'))

        # 4. 数据仪表盘
        if viz_path and os.path.exists(viz_path):
            doc.add_heading('二、 数据可视化概览', level=1)
            doc.add_picture(viz_path, width=Inches(6.0))
            p = doc.add_paragraph('图表 1: 价格对比、利润空间及众筹热度分析')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 5. 市场详情 (Amazon/AliExpress)
        doc.add_heading('三、 销售渠道详情 (Amazon / AliExpress)', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '平台'
        hdr_cells[1].text = '商品标题'
        hdr_cells[2].text = '价格'
        hdr_cells[3].text = '销量/评论'

        for item in sales_data[:10]: # 只展示前10条
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('platform', ''))
            row_cells[1].text = str(item.get('title', ''))[:50] + '...'
            row_cells[2].text = str(item.get('price', ''))
            row_cells[3].text = str(item.get('sold', item.get('reviews_count', 'N/A')))

        # 6. 供应链详情 (1688 / YiwuGo)
        doc.add_heading('四、 供应链详情 (1688 / 义乌购)', level=1)
        table_s = doc.add_table(rows=1, cols=3)
        table_s.style = 'Table Grid'
        hdr_cells = table_s.rows[0].cells
        hdr_cells[0].text = '平台'
        hdr_cells[1].text = '供应商/标题'
        hdr_cells[2].text = '进货价'

        for item in sourcing_data[:10]:
            row_cells = table_s.add_row().cells
            row_cells[0].text = str(item.get('platform', ''))
            row_cells[1].text = f"[{item.get('supplier', '未知')}] {item.get('title', '')[:40]}..."
            row_cells[2].text = str(item.get('price', ''))

        # 7. 创新趋势 (Kickstarter)
        if trend_data:
            doc.add_heading('五、 创新趋势参考 (Kickstarter)', level=1)
            for item in trend_data[:5]:
                doc.add_paragraph(f"项目: {item.get('title', '')}", style='List Bullet')
                doc.add_paragraph(f"  - 已筹金额: {item.get('pledged', '')} ({item.get('percent_funded', '')} funded)")
                doc.add_paragraph(f"  - 核心点: {item.get('description', '')}")

        # 保存文档
        safe_keyword = keyword.replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_path = os.path.join(self.output_dir, f"AnalysisReport_{safe_keyword}_{timestamp}.docx")
        doc.save(report_path)
        return report_path

